#!/usr/bin/env python3
# Genera las plantillas .docx de factura (renta anual + mensual) con diseño unificado.
# El diseño base es el de factura-anual-renta.html · la mensual lo reutiliza con
# las especificidades propias (subtítulo, columnas de la tabla, declaración).

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---- Paleta TWIM ----
GREEN_DARK = "173D30"
GREEN_MED  = "265C4B"
BEIGE      = "C2A78B"
CREAM      = "FDFCFA"
INK        = "2B2420"
MUTED      = "6B6258"
BORDER     = "E3DAC9"
FONT = "Century Gothic"

def shade_cell(cell, hex_color):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(sh)

def set_cell_borders(cell, color=BORDER, sz=4, sides=('top','bottom','left','right')):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def no_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for side, val in (('top',top),('bottom',bottom),('left',left),('right',right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)

def run(paragraph, text, size=11, color=INK, bold=False, italic=False,
        caps=False, spacing=None):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.bold = bold
    r.font.italic = italic
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(spacing))
        rPr.append(sp)
    return r

def para(container, align=None, space_before=0, space_after=4, line=1.3):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    return p

def build_invoice(path, *, titulo, subtitulo, servicios_headers, servicios_row,
                  sesiones_extra, declaracion):
    doc = Document()

    # --- Página A4, márgenes 0 (el contenido controla su padding) ---
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Mm(0)
    sec.bottom_margin = Mm(0)
    sec.left_margin = Mm(0)
    sec.right_margin = Mm(0)

    # ===== HEADER verde (tabla 1x2) =====
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    hl, hr = header.rows[0].cells
    hl.width = Mm(125)
    hr.width = Mm(85)
    for c in (hl, hr):
        shade_cell(c, GREEN_DARK)
        no_cell_borders(c)
        set_cell_margins(c, top=240, bottom=240, left=340, right=340)
        vcenter(c)

    # Header izquierda · logo + branding
    hl.paragraphs[0].clear()
    pic_p = hl.paragraphs[0]
    try:
        pic_p.add_run().add_picture('assets/logo-mindworld-on-dark.png', width=Mm(20))
    except Exception:
        pass
    pb = para(hl, space_after=0, line=1.25)
    run(pb, "TWIM PROJECT", size=9, color=BEIGE, bold=True, spacing=30)
    pb2 = para(hl, space_after=0, line=1.25)
    run(pb2, "Daniel Orozco Abia · CV11515", size=10, color=CREAM)

    # Header derecha · meta
    hr.paragraphs[0].clear()
    m1 = hr.paragraphs[0]
    m1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(m1, "Nº FACTURA  ", size=8.5, color=BEIGE, bold=True, spacing=14)
    run(m1, "{{NUMERO_FACTURA}}", size=10.5, color=CREAM)
    m2 = para(hr, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    run(m2, "FECHA DE EMISIÓN  ", size=8.5, color=BEIGE, bold=True, spacing=14)
    run(m2, "{{FECHA_EMISION}}", size=10.5, color=CREAM)

    # ===== CUERPO (en una tabla 1x1 con márgenes para simular padding 32px) =====
    body_tbl = doc.add_table(rows=1, cols=1)
    body_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    body = body_tbl.rows[0].cells[0]
    body.width = Mm(210)
    no_cell_borders(body)
    set_cell_margins(body, top=520, bottom=400, left=600, right=600)
    body.paragraphs[0].clear()

    # Título
    t = body.paragraphs[0]
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(3)
    run(t, titulo, size=21, color=GREEN_DARK, bold=True)

    # Subtítulo
    st = para(body, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    run(st, subtitulo, size=10.5, color=MUTED)

    # Regla beige
    rule = para(body, space_before=2, space_after=14)
    pPr = rule._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BEIGE)
    bd.append(bottom)
    pPr.append(bd)

    # ===== PARTES · emisor / receptor (tabla 1x2 con bordes beige) =====
    parties = body.add_table(rows=1, cols=2)
    parties.alignment = WD_TABLE_ALIGNMENT.CENTER
    pl, pr = parties.rows[0].cells
    pl.width = Mm(94)
    pr.width = Mm(94)

    def fill_party(cell, role, name, lines):
        for c in (cell,):
            set_cell_borders(c, color=BEIGE, sz=6)
            set_cell_margins(c, top=160, bottom=160, left=200, right=200)
        cell.paragraphs[0].clear()
        rp = cell.paragraphs[0]
        rp.paragraph_format.space_after = Pt(6)
        run(rp, role, size=8.5, color=GREEN_DARK, bold=True, spacing=22, caps=True)
        npar = para(cell, space_after=3)
        run(npar, name, size=12, color=GREEN_DARK, bold=True)
        for ln in lines:
            lp = para(cell, space_after=1, line=1.35)
            run(lp, ln, size=9.5, color=INK)

    fill_party(pl, "Emisor (profesional)", "Daniel Orozco Abia", [
        "NIF: 48442701D",
        "Psicólogo General Sanitario · Col. CV11515",
        "Calle Embajador Vich 3, 1B",
        "46002 Valencia · España",
        "Tel.: 625 231 297",
        "equipo@theworldismindproject.com",
    ])
    fill_party(pr, "Receptor (paciente)", "{{NOMBRE_PACIENTE}}", [
        "DNI: {{DNI_PACIENTE}}",
        "{{DIRECCION_PACIENTE_LINEA1}}",
        "{{DIRECCION_PACIENTE_LINEA2}}",
    ])

    para(body, space_after=6)  # espaciador

    # ===== TABLA DE SERVICIOS =====
    ncols = len(servicios_headers)
    serv = body.add_table(rows=2, cols=ncols)
    serv.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Cabecera
    for i, h in enumerate(servicios_headers):
        c = serv.rows[0].cells[i]
        shade_cell(c, GREEN_DARK)
        set_cell_borders(c, color=GREEN_DARK, sz=4)
        set_cell_margins(c, top=120, bottom=120, left=140, right=140)
        c.paragraphs[0].clear()
        hp = c.paragraphs[0]
        if h.get('right'):
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run(hp, h['t'], size=8.5, color=CREAM, bold=True, spacing=14, caps=True)
    # Fila de datos
    for i, val in enumerate(servicios_row):
        c = serv.rows[1].cells[i]
        set_cell_borders(c, color=BORDER, sz=4)
        set_cell_margins(c, top=150, bottom=150, left=140, right=140)
        c.paragraphs[0].clear()
        dp = c.paragraphs[0]
        right = servicios_headers[i].get('right')
        if right:
            dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run(dp, val, size=10.5, color=INK, bold=bool(right))

    # Fila extra opcional (sesiones del mes)
    if sesiones_extra:
        extra = body.add_table(rows=1, cols=1)
        extra.alignment = WD_TABLE_ALIGNMENT.CENTER
        ec = extra.rows[0].cells[0]
        shade_cell(ec, "F2F7F4")
        set_cell_borders(ec, color=BORDER, sz=4)
        set_cell_margins(ec, top=120, bottom=120, left=140, right=140)
        ec.paragraphs[0].clear()
        ep = ec.paragraphs[0]
        run(ep, sesiones_extra, size=9.5, color="5C6F66", italic=True)

    # Nota IVA
    iva = para(body, space_before=8, space_after=14)
    run(iva, "* Servicio exento de IVA según art. 20.1.3º Ley 37/1992 "
             "(servicios de asistencia sanitaria).", size=8.5, color=MUTED, italic=True)

    # ===== TOTAL BOX verde =====
    total = body.add_table(rows=1, cols=2)
    total.alignment = WD_TABLE_ALIGNMENT.CENTER
    tl, tr = total.rows[0].cells
    tl.width = Mm(120)
    tr.width = Mm(68)
    for c in (tl, tr):
        shade_cell(c, GREEN_DARK)
        no_cell_borders(c)
        set_cell_margins(c, top=170, bottom=170, left=240, right=240)
        vcenter(c)
    tl.paragraphs[0].clear()
    tlp = tl.paragraphs[0]
    run(tlp, "Total a abonar", size=11, color=CREAM, bold=True, spacing=14, caps=True)
    tr.paragraphs[0].clear()
    trp = tr.paragraphs[0]
    trp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(trp, "{{IMPORTE}}", size=18, color=CREAM, bold=True)

    para(body, space_after=10)

    # Declaración
    for ln in declaracion:
        dp = para(body, space_after=2, line=1.5)
        run(dp, ln, size=10, color=INK)

    # ===== FIRMA =====
    para(body, space_after=24)
    sig_line = para(body, space_after=4)
    pPr = sig_line._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), INK)
    bd.append(bottom)
    pPr.append(bd)
    run(sig_line, " ", size=10)
    sn = para(body, space_after=1)
    run(sn, "Daniel Orozco Abia", size=11, color=GREEN_DARK, bold=True)
    sm = para(body, space_after=0)
    run(sm, "Psicólogo General Sanitario · Col. CV11515", size=9.5, color=MUTED)

    # ===== FOOTER =====
    foot = para(body, space_before=22, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    pPr = foot._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), BEIGE)
    bd.append(top)
    pPr.append(bd)
    run(foot, "Calle Embajador Vich 3, 1B  ·  46002 Valencia  ·  "
              "Tel. 625 231 297  ·  equipo@theworldismindproject.com",
        size=9, color=MUTED)

    doc.save(path)
    print(f"Generado · {path}")


# ===== FACTURA ANUAL · RENTA =====
build_invoice(
    'documentos-internos/plantillas/factura-anual-renta.docx',
    titulo="FACTURA DE HONORARIOS",
    subtitulo="Ejercicio fiscal {{ANIO}} · Honorarios por tratamiento psicológico",
    servicios_headers=[
        {'t': 'Descripción del servicio'},
        {'t': 'Período'},
        {'t': 'Sesiones'},
        {'t': 'Importe', 'right': True},
    ],
    servicios_row=[
        "Honorarios por tratamiento psicológico",
        "{{PERIODO}}",
        "{{NUM_SESIONES}}",
        "{{IMPORTE}}",
    ],
    sesiones_extra=None,
    declaracion=[
        "El presente documento acredita los honorarios satisfechos en concepto de "
        "tratamiento psicológico durante el ejercicio fiscal {{ANIO}}, a efectos de "
        "declaración del IRPF.",
        "Emitido en Valencia, a {{FECHA_EMISION}}.",
    ],
)

# ===== FACTURA MENSUAL · PACIENTE (mismo diseño, especificidades propias) =====
build_invoice(
    'documentos-internos/plantillas/factura-mensual-paciente.docx',
    titulo="FACTURA DE HONORARIOS",
    subtitulo="{{MES}} · Honorarios por tratamiento psicológico",
    servicios_headers=[
        {'t': 'Descripción del servicio'},
        {'t': 'Sesiones'},
        {'t': 'Importe', 'right': True},
    ],
    servicios_row=[
        "Honorarios por tratamiento psicológico",
        "{{NUM_SESIONES}}",
        "{{IMPORTE}}",
    ],
    sesiones_extra="Sesiones {{MES}}: {{SESIONES_FECHAS}}",
    declaracion=[
        "El presente documento acredita los honorarios satisfechos en concepto de "
        "tratamiento psicológico correspondientes a {{MES}}.",
        "Emitido en Valencia, a {{FECHA_EMISION}}.",
    ],
)
print("OK · 2 plantillas .docx generadas")
PY_END_MARKER_UNUSED = None
